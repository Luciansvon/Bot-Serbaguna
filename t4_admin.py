import os
import json
import logging
import httpx
from datetime import datetime
from pathlib import Path
from crewai import Agent
from crewai.tools import BaseTool
from config import admin_llm, OUTPUT_DIR

logger = logging.getLogger('bima_core')

# ============================================================
# STYLE PRESETS — 5 gaya tulisan dengan palet warna sendiri
# ============================================================
STYLES = {
    "formal": {
        "title_rgb": (43, 74, 138),
        "accent_rgb": (43, 74, 138),
        "table_header_rgb": (43, 74, 138),
        "table_alt_rgb": (238, 242, 255),
        "title_size": 18,
        "body_size": 10,
        "heading_size": 13,
        "tone": "Profesional, baku, struktur kaku, kalimat formal.",
        "label": "FORMAL",
        "font_family": "Calibri",
        "pdf_font": "Helvetica",
        "line_spacing": 1.15,
        "justify": False,
        "margins_cm": {"top": 2.54, "bottom": 2.54, "left": 2.54, "right": 2.54},
    },
    "casual": {
        "title_rgb": (245, 158, 11),
        "accent_rgb": (217, 119, 6),
        "table_header_rgb": (245, 158, 11),
        "table_alt_rgb": (255, 247, 237),
        "title_size": 20,
        "body_size": 11,
        "heading_size": 14,
        "tone": "Hangat, conversational, ramah, kalimat lebih luwes.",
        "label": "CASUAL",
        "font_family": "Calibri",
        "pdf_font": "Helvetica",
        "line_spacing": 1.15,
        "justify": False,
        "margins_cm": {"top": 2.54, "bottom": 2.54, "left": 2.54, "right": 2.54},
    },
    "creative": {
        "title_rgb": (168, 85, 247),
        "accent_rgb": (236, 72, 153),
        "table_header_rgb": (168, 85, 247),
        "table_alt_rgb": (250, 245, 255),
        "title_size": 22,
        "body_size": 11,
        "heading_size": 15,
        "tone": "Ekspresif, naratif, deskriptif, gunakan metafora secukupnya.",
        "label": "CREATIVE",
        "font_family": "Calibri",
        "pdf_font": "Helvetica",
        "line_spacing": 1.15,
        "justify": False,
        "margins_cm": {"top": 2.54, "bottom": 2.54, "left": 2.54, "right": 2.54},
    },
    "technical": {
        "title_rgb": (31, 41, 55),
        "accent_rgb": (75, 85, 99),
        "table_header_rgb": (31, 41, 55),
        "table_alt_rgb": (243, 244, 246),
        "title_size": 16,
        "body_size": 10,
        "heading_size": 12,
        "tone": "Presisi, padat, terminologi teknis OK, struktur step-by-step.",
        "label": "TECHNICAL",
        "font_family": "Calibri",
        "pdf_font": "Helvetica",
        "line_spacing": 1.15,
        "justify": False,
        "margins_cm": {"top": 2.54, "bottom": 2.54, "left": 2.54, "right": 2.54},
    },
    "educational": {
        "title_rgb": (16, 185, 129),
        "accent_rgb": (5, 150, 105),
        "table_header_rgb": (16, 185, 129),
        "table_alt_rgb": (236, 253, 245),
        "title_size": 18,
        "body_size": 11,
        "heading_size": 13,
        "tone": "Step-by-step, jelas, ramah pemula, pakai contoh konkret.",
        "label": "EDUCATIONAL",
        "font_family": "Calibri",
        "pdf_font": "Helvetica",
        "line_spacing": 1.15,
        "justify": False,
        "margins_cm": {"top": 2.54, "bottom": 2.54, "left": 2.54, "right": 2.54},
    },
    "academic": {
        "title_rgb": (0, 0, 0),
        "accent_rgb": (0, 0, 0),
        "table_header_rgb": (50, 50, 50),
        "table_alt_rgb": (240, 240, 240),
        "title_size": 14,
        "body_size": 12,
        "heading_size": 12,
        "tone": "Sangat formal, objektif, impersonal (tidak pakai saya/kami), berdasarkan fakta dan rujukan ilmiah, bahasa baku.",
        "label": "ACADEMIC",
        "font_family": "Times New Roman",
        "pdf_font": "Times",
        "line_spacing": 1.5,
        "justify": True,
        "margins_cm": {"top": 4, "bottom": 3, "left": 4, "right": 3},
    },
}


def detect_style(text: str) -> str:
    """Auto-detect gaya tulisan dari permintaan user."""
    t = (text or "").lower()
    if any(k in t for k in ["skripsi", "tugas akhir", "tesis", "disertasi", "jurnal ilmiah",
                             "akademik", "academic", "thesis", "makalah", "karya ilmiah"]):
        return "academic"
    if any(k in t for k in ["formal", "resmi", "profesional", "kontrak", "proposal", "baku", "perusahaan", "bisnis"]):
        return "formal"
    if any(k in t for k in ["santai", "casual", "blog", "newsletter", "ngobrol", "ramah", "asik", "ringan"]):
        return "casual"
    if any(k in t for k in ["kreatif", "naratif", "cerita", "story", "esai", "puisi", "creative", "novel"]):
        return "creative"
    if any(k in t for k in ["teknis", "technical", "dokumentasi", "spek", "api", "manual", "code", "kode"]):
        return "technical"
    if any(k in t for k in ["tutorial", "panduan", "lesson", "edukasi", "ajar", "pemula", "step", "belajar"]):
        return "educational"
    return "formal"


def detect_format(text: str) -> str:
    """Auto-detect format dokumen yang diminta (pdf / word / excel)."""
    t = (text or "").lower()
    if any(k in t for k in ["excel", "xlsx", "spreadsheet", "tabel data", "rekap"]):
        return "excel"
    if any(k in t for k in [".docx", "word", "dokumen word"]):
        return "word"
    return "pdf"  # default


def _hex_from_rgb(rgb: tuple) -> str:
    return "{:02X}{:02X}{:02X}".format(*rgb)


def _render_chart(chart: dict, style: dict) -> str:
    """Render chart spec ke PNG, simpan ke OUTPUT_DIR, return path lokal.
    Format chart sama seperti Chart.js: type, title, labels, datasets[{label, data}].
    Type yang didukung: bar, line, pie."""
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt

    chart_type = chart.get("type", "bar")
    title = chart.get("title", "")
    labels = chart.get("labels", [])
    datasets = chart.get("datasets", [])
    if not datasets:
        raise ValueError("Chart datasets kosong")

    primary_rgb = tuple(c / 255 for c in style["accent_rgb"])
    title_rgb = tuple(c / 255 for c in style["title_rgb"])
    palette = [
        primary_rgb,
        tuple(c / 255 for c in style["table_header_rgb"]),
        (0.40, 0.55, 0.85),
        (0.85, 0.55, 0.40),
        (0.55, 0.85, 0.55),
        (0.85, 0.45, 0.55),
    ]

    fig, ax = plt.subplots(figsize=(8, 4.5), dpi=120)

    if chart_type == "bar":
        n_ds = len(datasets)
        x = list(range(len(labels)))
        width = 0.8 / max(n_ds, 1)
        for i, ds in enumerate(datasets):
            offsets = [xi + width * i for xi in x]
            ax.bar(offsets, ds["data"], width=width * 0.9,
                   label=ds.get("label", f"Series {i + 1}"),
                   color=palette[i % len(palette)])
        ax.set_xticks([xi + width * (n_ds - 1) / 2 for xi in x])
        ax.set_xticklabels(labels, rotation=30, ha='right', fontsize=9)
    elif chart_type == "line":
        for i, ds in enumerate(datasets):
            ax.plot(labels, ds["data"], marker='o', linewidth=2,
                    label=ds.get("label", f"Series {i + 1}"),
                    color=palette[i % len(palette)])
        ax.tick_params(axis='x', rotation=30)
        ax.grid(True, alpha=0.3)
    elif chart_type == "pie":
        ax.pie(datasets[0]["data"], labels=labels, autopct='%1.1f%%',
               startangle=140, colors=palette[:len(labels)])
        ax.axis('equal')
    else:
        raise ValueError(f"Tipe chart tidak didukung: {chart_type}")

    if title:
        ax.set_title(title, fontsize=13, color=title_rgb, fontweight='bold', pad=12)
    if chart_type != "pie" and len(datasets) > 1:
        ax.legend(loc='best', fontsize=9, framealpha=0.9)

    plt.tight_layout()
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S_%f')
    path = OUTPUT_DIR / f"chart_{chart_type}_{timestamp}.png"
    plt.savefig(path, dpi=120, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    return str(path)


# ============================================================
# Excel Generator — sekarang support formula + style
# ============================================================
class ExcelGeneratorTool(BaseTool):
    name: str = "Excel Generator Tool"
    description: str = """Buat file Excel (.xlsx) dari data JSON.
    Input format JSON string:
    {
        "filename": "nama_file",
        "style": "formal" | "casual" | "creative" | "technical" | "educational",
        "sheets": [
            {
                "name": "Sheet1",
                "headers": ["Kolom1", "Kolom2", "Kolom3"],
                "rows": [
                    ["data1", "data2", 100],
                    ["data4", "data5", "=SUM(C2:C2)"]
                ]
            }
        ],
        "references": [
            {"text": "Penulis (2026). Judul.", "url": "https://..."}
        ]
    }
    Cell value yang diawali '=' akan otomatis jadi formula Excel.
    Field 'references' (opsional): auto-bikin sheet "Referensi" dengan kolom No|Sumber|URL (URL clickable).
    URL referensi WAJIB valid (Wikipedia/situs resmi/jurnal open-access). JANGAN dikarang.
    Field 'style' opsional — default 'formal'."""

    def _run(self, input_json: str) -> str:
        try:
            import openpyxl
            from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

            try:
                data = json.loads(input_json)
            except json.JSONDecodeError as e:
                return f"FAILED|JSON tidak valid: {e}"

            style_name = data.get("style", "formal")
            style = STYLES.get(style_name, STYLES["formal"])
            header_hex = _hex_from_rgb(style["table_header_rgb"])
            alt_hex = _hex_from_rgb(style["table_alt_rgb"])
            accent_hex = _hex_from_rgb(style["accent_rgb"])

            wb = openpyxl.Workbook()
            wb.remove(wb.active)

            header_fill = PatternFill(start_color=header_hex, end_color=header_hex, fill_type="solid")
            alt_fill = PatternFill(start_color=alt_hex, end_color=alt_hex, fill_type="solid")
            header_font = Font(bold=True, color="FFFFFF", size=11)
            total_font = Font(bold=True, size=11)
            border = Border(
                left=Side(style='thin'), right=Side(style='thin'),
                top=Side(style='thin'), bottom=Side(style='thin')
            )

            # === SUMMARY SHEET — best practice: sheet pertama sebagai ringkasan ===
            sheets_data = data.get("sheets", [])
            doc_title = data.get("title") or data.get("filename", "Dokumen")
            sum_ws = wb.create_sheet(title="Ringkasan")
            sum_ws.merge_cells(start_row=1, start_column=1, end_row=1,
                               end_column=max(len(sheets_data), 4))
            tc = sum_ws.cell(row=1, column=1, value=str(doc_title).upper())
            tc.font = Font(bold=True, size=16, color="FFFFFF")
            tc.fill = PatternFill(start_color=header_hex, end_color=header_hex, fill_type="solid")
            tc.alignment = Alignment(horizontal='center', vertical='center')
            sum_ws.row_dimensions[1].height = 42
            for i, (k, v) in enumerate([
                ("Dokumen", str(doc_title)),
                ("Tanggal Dibuat", datetime.now().strftime("%d %B %Y, %H:%M")),
                ("Style", style["label"]),
                ("Jumlah Sheet Data", len(sheets_data)),
            ], 3):
                sum_ws.cell(row=i, column=1, value=k).font = Font(bold=True, size=10)
                sum_ws.cell(row=i, column=1).alignment = Alignment(indent=1)
                sum_ws.cell(row=i, column=2, value=v).alignment = Alignment(indent=1)
            sum_ws.column_dimensions['A'].width = 22
            sum_ws.column_dimensions['B'].width = 42

            for sheet_data in sheets_data:
                ws = wb.create_sheet(title=sheet_data.get("name", "Sheet")[:31])
                headers = sheet_data.get("headers", [])
                rows = sheet_data.get("rows", [])
                total_row_indices = set(sheet_data.get("total_rows", []))

                # Judul sheet opsional di baris pertama
                data_start_row = 1
                sheet_title = sheet_data.get("title", "")
                if sheet_title:
                    ws.merge_cells(start_row=1, start_column=1,
                                   end_row=1, end_column=max(len(headers), 1))
                    st_cell = ws.cell(row=1, column=1, value=sheet_title)
                    st_cell.font = Font(bold=True, size=13,
                                        color=_hex_from_rgb(style["title_rgb"]))
                    st_cell.alignment = Alignment(horizontal='left', indent=1)
                    ws.row_dimensions[1].height = 28
                    data_start_row = 2

                hdr_row = data_start_row
                for col_idx, header in enumerate(headers, 1):
                    cell = ws.cell(row=hdr_row, column=col_idx, value=header)
                    cell.font = header_font
                    cell.fill = header_fill
                    cell.alignment = Alignment(horizontal='center', vertical='center',
                                              wrap_text=True, indent=1)
                    cell.border = border

                for local_idx, row_data in enumerate(rows):
                    row_idx = hdr_row + 1 + local_idx
                    is_total = local_idx in total_row_indices
                    for col_idx, value in enumerate(row_data, 1):
                        cell = ws.cell(row=row_idx, column=col_idx, value=value)
                        if is_total:
                            cell.border = Border(top=Side(style='medium'),
                                                 bottom=Side(style='double'))
                            cell.font = total_font
                            cell.fill = PatternFill(start_color=alt_hex, end_color=alt_hex,
                                                    fill_type="solid")
                        else:
                            cell.border = border
                            if row_idx % 2 == 0:
                                cell.fill = alt_fill

                        # Smart alignment — numerik kanan, teks kiri (standar finansial)
                        if isinstance(value, (int, float)) and not str(value).startswith('='):
                            cell.alignment = Alignment(horizontal='right', vertical='center',
                                                       indent=1)
                            if isinstance(value, float) and value != int(value):
                                cell.number_format = '#,##0.00'
                            else:
                                cell.number_format = '#,##0'
                        else:
                            cell.alignment = Alignment(horizontal='left', vertical='center',
                                                       wrap_text=True, indent=1)

                for col in ws.columns:
                    max_len = max((len(str(c.value or "")) for c in col), default=10)
                    ws.column_dimensions[col[0].column_letter].width = max(14, min(max_len + 6, 50))

                ws.row_dimensions[hdr_row].height = 30
                if headers:
                    ws.freeze_panes = ws.cell(row=hdr_row + 1, column=1).coordinate

            # === Sheet Referensi (opsional) ===
            references = data.get("references", [])
            if references:
                from openpyxl.styles import Font as _Font
                ref_ws = wb.create_sheet(title="Referensi")
                ref_headers = ["No", "Sumber", "URL"]
                for col_idx, h in enumerate(ref_headers, 1):
                    cell = ref_ws.cell(row=1, column=col_idx, value=h)
                    cell.font = header_font
                    cell.fill = header_fill
                    cell.alignment = Alignment(horizontal='center', vertical='center', indent=1)
                    cell.border = border

                accent_hex = _hex_from_rgb(style["accent_rgb"])
                link_font = _Font(color=accent_hex, underline="single")

                for idx, ref in enumerate(references, 1):
                    if isinstance(ref, dict):
                        ref_text = ref.get("text", "")
                        ref_url = ref.get("url", "")
                    else:
                        ref_text = str(ref)
                        ref_url = ""

                    no_cell = ref_ws.cell(row=idx + 1, column=1, value=idx)
                    no_cell.alignment = Alignment(horizontal='center', vertical='center')
                    no_cell.border = border

                    src_cell = ref_ws.cell(row=idx + 1, column=2, value=ref_text)
                    src_cell.alignment = Alignment(vertical='center', wrap_text=True, indent=1)
                    src_cell.border = border

                    url_cell = ref_ws.cell(row=idx + 1, column=3, value=ref_url)
                    if ref_url:
                        url_cell.hyperlink = ref_url
                        url_cell.font = link_font
                    url_cell.alignment = Alignment(vertical='center', wrap_text=True, indent=1)
                    url_cell.border = border

                    if (idx + 1) % 2 == 0:
                        no_cell.fill = alt_fill
                        src_cell.fill = alt_fill
                        # url_cell font sudah punya color khusus, tetap kasih fill
                        url_cell.fill = alt_fill

                ref_ws.column_dimensions['A'].width = 6
                ref_ws.column_dimensions['B'].width = 60
                ref_ws.column_dimensions['C'].width = 50
                ref_ws.row_dimensions[1].height = 30
                ref_ws.freeze_panes = "A2"

            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{data.get('filename', 'laporan')}_{timestamp}.xlsx"
            filepath = OUTPUT_DIR / filename
            wb.save(filepath)

            return f"SUCCESS|{filepath}|Excel ({style['label']}) berhasil dibuat: {filename}"
        except Exception as e:
            logger.error(f"[ADMIN] ExcelGenerator error: {e}", exc_info=True)
            return f"FAILED|{e}"


# ============================================================
# Word Generator — sekarang support style + image embedding
# ============================================================
class WordGeneratorTool(BaseTool):
    name: str = "Word Generator Tool"
    description: str = """Buat file Word (.docx) dengan style fleksibel.
    Input format JSON string:
    {
        "filename": "nama_file",
        "style": "formal" | "casual" | "creative" | "technical" | "educational",
        "title": "Judul Dokumen",
        "subtitle": "Subjudul opsional",
        "author": "Nama Author",
        "toc": true,
        "sections": [
            {
                "heading": "Judul Bagian",
                "content": "Isi paragraf...",
                "list": ["Item 1", "Item 2"],
                "key_values": {"Nama": "Bima", "Jabatan": "Admin"},
                "image_path": "/path/to/image.png",
                "charts": [
                    {"type": "bar", "title": "Penjualan 2026", "labels": ["Q1","Q2","Q3"],
                     "datasets": [{"label": "Unit", "data": [100, 150, 200]}]}
                ],
                "table": {
                    "headers": ["Kolom1", "Kolom2"],
                    "rows": [["data1", "data2"]]
                }
            }
        ],
        "references": [
            {"text": "Penulis (2026). Judul paper. Penerbit.",
             "url": "https://doi.org/..."}
        ]
    }
    Field 'charts' (opsional, list): tipe 'bar'|'line'|'pie' — auto-render via matplotlib pakai warna dari style.
    Field 'references' (opsional): daftar pustaka — auto render section terakhir dengan hyperlink clickable.
    URL referensi WAJIB valid (Wikipedia/situs resmi/jurnal open-access). JANGAN dikarang.
    Tipe dokumen support: laporan, proposal, surat, resume, blog, tutorial — gaya menyesuaikan field 'style'."""

    def _run(self, input_json: str) -> str:
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import nsdecls
            from docx.oxml import parse_xml

            def set_cell_bg(cell, hex_color: str):
                shd = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{hex_color}"/>')
                cell._tc.get_or_add_tcPr().append(shd)

            def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
                """Atur padding cell (twips, 1 inch = 1440 twips, ~80 = 0.055 inch)."""
                from docx.oxml import OxmlElement
                from docx.oxml.ns import qn
                tc_pr = cell._tc.get_or_add_tcPr()
                tc_mar = OxmlElement('w:tcMar')
                for direction, val in (('top', top), ('bottom', bottom), ('left', left), ('right', right)):
                    node = OxmlElement(f'w:{direction}')
                    node.set(qn('w:w'), str(val))
                    node.set(qn('w:type'), 'dxa')
                    tc_mar.append(node)
                tc_pr.append(tc_mar)

            def add_hyperlink(paragraph, text: str, url: str, color_hex: str = "0563C1"):
                """Tambah hyperlink clickable ke paragraph (python-docx tidak punya API langsung)."""
                from docx.oxml import OxmlElement
                from docx.oxml.ns import qn
                part = paragraph.part
                r_id = part.relate_to(
                    url,
                    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                    is_external=True,
                )
                hyperlink = OxmlElement('w:hyperlink')
                hyperlink.set(qn('r:id'), r_id)
                new_run = OxmlElement('w:r')
                rPr = OxmlElement('w:rPr')
                color_node = OxmlElement('w:color')
                color_node.set(qn('w:val'), color_hex)
                rPr.append(color_node)
                u_node = OxmlElement('w:u')
                u_node.set(qn('w:val'), 'single')
                rPr.append(u_node)
                new_run.append(rPr)
                t = OxmlElement('w:t')
                t.text = text
                t.set(qn('xml:space'), 'preserve')
                new_run.append(t)
                hyperlink.append(new_run)
                paragraph._p.append(hyperlink)
                return hyperlink

            try:
                data = json.loads(input_json)
            except json.JSONDecodeError as e:
                return f"FAILED|JSON tidak valid: {e}"

            style_name = data.get("style", "formal")
            style = STYLES.get(style_name, STYLES["formal"])
            title_color = RGBColor(*style["title_rgb"])
            accent_color = RGBColor(*style["accent_rgb"])
            header_hex = _hex_from_rgb(style["table_header_rgb"])

            # === Configurable typography (JSON overrides > style defaults) ===
            font_family = data.get("font_family", style.get("font_family", "Calibri"))
            margins_cm = data.get("margins", style.get("margins_cm", {"top": 2.54, "bottom": 2.54, "left": 2.54, "right": 2.54}))
            do_justify = data.get("justify", style.get("justify", False))
            line_spacing_mult = data.get("line_spacing", style.get("line_spacing", 1.15))

            doc = Document()

            # === MARGIN — configurable via JSON atau style preset ===
            from docx.shared import Cm
            from docx.oxml import OxmlElement as _OxmlEl
            from docx.oxml.ns import qn as _qn
            section_doc = doc.sections[0]
            section_doc.top_margin    = Cm(margins_cm.get("top", 2.54))
            section_doc.bottom_margin = Cm(margins_cm.get("bottom", 2.54))
            section_doc.left_margin   = Cm(margins_cm.get("left", 2.54))
            section_doc.right_margin  = Cm(margins_cm.get("right", 2.54))

            # === GAYA NORMAL: font, line spacing, space after ===
            normal_style = doc.styles['Normal']
            normal_style.font.name = font_family
            normal_style.font.size = Pt(style["body_size"])
            normal_style.paragraph_format.line_spacing = Pt(style["body_size"] * line_spacing_mult)
            normal_style.paragraph_format.space_after = Pt(6)

            # === HEADING STYLES (multi-level support) ===
            _heading_configs = [
                (1, style["heading_size"] + 2, True, False),   # Level 1: besar, bold
                (2, style["heading_size"], True, False),       # Level 2: sedang, bold
                (3, style["heading_size"] - 1, True, True),    # Level 3: kecil, bold italic
            ]
            for _lvl, _sz, _bold, _italic in _heading_configs:
                _hstyle_name = f'Heading {_lvl}'
                if _hstyle_name in doc.styles:
                    _hs = doc.styles[_hstyle_name]
                    _hs.font.name = font_family
                    _hs.font.size = Pt(_sz)
                    _hs.font.bold = _bold
                    _hs.font.italic = _italic
                    _hs.font.color.rgb = title_color
                    _hs.paragraph_format.space_before = Pt(12)
                    _hs.paragraph_format.space_after = Pt(6)

            # === HELPER: Roman numeral converter ===
            def _to_roman(num):
                vals = [(1000,'m'),(900,'cm'),(500,'d'),(400,'cd'),(100,'c'),
                        (90,'xc'),(50,'l'),(40,'xl'),(10,'x'),(9,'ix'),(5,'v'),(4,'iv'),(1,'i')]
                result = ''
                for v, r in vals:
                    while num >= v:
                        result += r
                        num -= v
                return result

            # === HELPER: Add section break + set page number format ===
            def _set_page_num_format(section_obj, fmt='decimal', start=None):
                """Set page number format. fmt: 'decimal', 'lowerRoman', 'upperRoman'."""
                from docx.oxml import OxmlElement as _OE
                from docx.oxml.ns import qn as _qn
                sectPr = section_obj._sectPr
                pgNumType = _OE('w:pgNumType')
                pgNumType.set(_qn('w:fmt'), fmt)
                if start is not None:
                    pgNumType.set(_qn('w:start'), str(start))
                sectPr.append(pgNumType)

            # === PAGE NUMBERS DI FOOTER ===
            footer = section_doc.footer
            footer_para = footer.paragraphs[0]
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer_run = footer_para.add_run()
            footer_run.font.size = Pt(9)
            footer_run.font.color.rgb = RGBColor(0x90, 0x90, 0x90)
            # Tambah field PAGE dan NUMPAGES via XML
            def _add_fld(para, instr):
                from docx.oxml import OxmlElement as _OE
                from docx.oxml.ns import qn as _qn
                fldChar_b = _OE('w:fldChar'); fldChar_b.set(_qn('w:fldCharType'), 'begin')
                instrText = _OE('w:instrText'); instrText.text = instr
                instrText.set(_qn('xml:space'), 'preserve')
                fldChar_e = _OE('w:fldChar'); fldChar_e.set(_qn('w:fldCharType'), 'end')
                r = _OE('w:r')
                r.append(fldChar_b)
                p = _OE('w:r')
                p.append(instrText)
                e = _OE('w:r')
                e.append(fldChar_e)
                para._p.extend([r, p, e])
            footer_para.add_run(f"{data.get('author', 'B.I.M.A Core')}  |  halaman ")
            _add_fld(footer_para, ' PAGE ')
            footer_para.add_run(" dari ")
            _add_fld(footer_para, ' NUMPAGES ')

            # === Set front matter to Roman numerals ===
            is_academic = (style_name == 'academic')
            if is_academic:
                _set_page_num_format(section_doc, fmt='lowerRoman', start=1)

            # Title
            title_para = doc.add_paragraph()
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_para.add_run(data.get("title", "Dokumen"))
            title_run.font.bold = True
            title_run.font.size = Pt(style["title_size"] + 4)
            title_run.font.color.rgb = title_color

            # Subtitle (opsional)
            if data.get("subtitle"):
                sub_para = doc.add_paragraph()
                sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                sub_run = sub_para.add_run(data["subtitle"])
                sub_run.font.italic = True
                sub_run.font.size = Pt(style["body_size"] + 2)
                sub_run.font.color.rgb = accent_color

            # Author + tanggal
            info = doc.add_paragraph()
            info.alignment = WD_ALIGN_PARAGRAPH.CENTER
            info_run = info.add_run(f"{data.get('author', 'B.I.M.A Core')} · {datetime.now().strftime('%d %B %Y')} · [{style['label']}]")
            info_run.font.color.rgb = RGBColor(0x70, 0x70, 0x70)
            info_run.font.size = Pt(9)
            doc.add_paragraph()

            # === ABSTRAK (opsional — khusus dokumen akademik) ===
            if data.get("abstract"):
                doc.add_page_break()
                abs_title = doc.add_paragraph()
                abs_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                abs_title_run = abs_title.add_run("ABSTRAK")
                abs_title_run.font.bold = True
                abs_title_run.font.size = Pt(style["heading_size"] + 2)
                abs_title_run.font.color.rgb = title_color
                abs_title_run.font.name = font_family

                abs_para = doc.add_paragraph(data["abstract"])
                abs_para.paragraph_format.space_after = Pt(8)
                # Abstrak selalu single-spaced meskipun body 1.5
                abs_para.paragraph_format.line_spacing = Pt(style["body_size"] * 1.15)
                abs_para.paragraph_format.left_indent = Inches(0.5)
                abs_para.paragraph_format.right_indent = Inches(0.5)
                if do_justify:
                    abs_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                # Keywords
                if data.get("keywords"):
                    kw_para = doc.add_paragraph()
                    kw_para.paragraph_format.left_indent = Inches(0.5)
                    kw_label = kw_para.add_run("Kata Kunci: ")
                    kw_label.font.bold = True
                    kw_label.font.size = Pt(style["body_size"])
                    kw_label.font.name = font_family
                    kw_value = kw_para.add_run(", ".join(data["keywords"]))
                    kw_value.font.italic = True
                    kw_value.font.size = Pt(style["body_size"])
                    kw_value.font.name = font_family

            # === Table of Contents (opsional, multi-level) ===
            sections = data.get("sections", [])
            if data.get("toc") and sections:
                doc.add_page_break()
                toc_h = doc.add_paragraph()
                toc_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
                toc_run = toc_h.add_run("DAFTAR ISI")
                toc_run.font.bold = True
                toc_run.font.size = Pt(style["heading_size"] + 2)
                toc_run.font.color.rgb = title_color
                toc_run.font.name = font_family
                doc.add_paragraph()

                for sec in sections:
                    if sec.get("heading"):
                        lvl = sec.get("level", 1)
                        toc_p = doc.add_paragraph()
                        toc_p.paragraph_format.space_after = Pt(2)
                        # Indentasi berdasarkan level
                        toc_p.paragraph_format.left_indent = Inches(0.3 * (lvl - 1))
                        toc_run = toc_p.add_run(sec["heading"])
                        toc_run.font.size = Pt(style["body_size"])
                        toc_run.font.name = font_family
                        if lvl == 1:
                            toc_run.font.bold = True

            # === SECTION BREAK: Roman → Arabic page numbering ===
            if is_academic:
                from docx.oxml import OxmlElement as _OE_sb
                from docx.oxml.ns import qn as _qn_sb
                # Add section break (new page) before body content
                body_break_para = doc.add_paragraph()
                pPr = body_break_para._p.get_or_add_pPr()
                sectPr = _OE_sb('w:sectPr')
                sectType = _OE_sb('w:type')
                sectType.set(_qn_sb('w:val'), 'nextPage')
                sectPr.append(sectType)
                # Copy margins to new section
                pgMar = _OE_sb('w:pgMar')
                pgMar.set(_qn_sb('w:top'), str(int(margins_cm.get('top', 2.54) * 567)))
                pgMar.set(_qn_sb('w:bottom'), str(int(margins_cm.get('bottom', 2.54) * 567)))
                pgMar.set(_qn_sb('w:left'), str(int(margins_cm.get('left', 2.54) * 567)))
                pgMar.set(_qn_sb('w:right'), str(int(margins_cm.get('right', 2.54) * 567)))
                sectPr.append(pgMar)
                # Set Arabic numbering starting from 1
                pgNumType = _OE_sb('w:pgNumType')
                pgNumType.set(_qn_sb('w:fmt'), 'decimal')
                pgNumType.set(_qn_sb('w:start'), '1')
                sectPr.append(pgNumType)
                pPr.append(sectPr)
            else:
                if data.get("toc") and sections:
                    doc.add_page_break()

            for section in sections:
                if section.get("heading"):
                    lvl = section.get("level", 1)
                    lvl = max(1, min(lvl, 3))  # clamp 1-3

                    # Gunakan built-in Heading style (sudah dikustomisasi di atas)
                    h = doc.add_heading(section["heading"], level=lvl)
                    # Override font untuk memastikan konsistensi
                    for run in h.runs:
                        run.font.name = font_family
                        run.font.color.rgb = title_color

                    # Garis aksen hanya untuk heading level 1
                    if lvl == 1:
                        h.paragraph_format.space_after = Pt(4)
                        from docx.oxml import OxmlElement as _OE2
                        from docx.oxml.ns import qn as _qn2
                        pPr = h._p.get_or_add_pPr()
                        pBdr = _OE2('w:pBdr')
                        bottom = _OE2('w:bottom')
                        bottom.set(_qn2('w:val'), 'single')
                        bottom.set(_qn2('w:sz'), '4')
                        bottom.set(_qn2('w:space'), '4')
                        bottom.set(_qn2('w:color'), _hex_from_rgb(style["accent_rgb"]))
                        pBdr.append(bottom)
                        pPr.append(pBdr)
                    elif lvl == 2:
                        h.paragraph_format.space_after = Pt(4)
                        h.paragraph_format.left_indent = Inches(0.15)
                    elif lvl == 3:
                        h.paragraph_format.space_after = Pt(3)
                        h.paragraph_format.left_indent = Inches(0.3)

                if section.get("content"):
                    p = doc.add_paragraph(section["content"])
                    p.paragraph_format.space_after = Pt(8)
                    p.paragraph_format.line_spacing = Pt(style["body_size"] * line_spacing_mult)
                    if do_justify:
                        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                # Image embedding
                if section.get("image_path"):
                    img_path = Path(section["image_path"])
                    if img_path.exists() and img_path.is_file():
                        try:
                            doc.add_picture(str(img_path), width=Inches(5.5))
                        except Exception as img_err:
                            logger.warning(f"[ADMIN] Gagal embed image {img_path}: {img_err}")

                # Chart rendering (matplotlib → embed PNG)
                for chart in section.get("charts", []) or []:
                    try:
                        chart_path = _render_chart(chart, style)
                        doc.add_picture(chart_path, width=Inches(6))
                    except Exception as chart_err:
                        logger.warning(f"[ADMIN] Gagal render chart Word: {chart_err}")

                # Bullet list
                if section.get("list"):
                    for item in section["list"]:
                        doc.add_paragraph(str(item), style='List Bullet')

                # Key Values (untuk surat izin, detail rapi dengan titik dua sejajar)
                if section.get("key_values") and isinstance(section["key_values"], dict):
                    kv_table = doc.add_table(rows=len(section["key_values"]), cols=3)
                    for idx, (k, v) in enumerate(section["key_values"].items()):
                        row = kv_table.rows[idx].cells
                        row[0].text = str(k)
                        row[0].width = Inches(1.5)
                        row[1].text = ":"
                        row[1].width = Inches(0.2)
                        row[2].text = str(v)
                        # Remove borders for clean look by leaving as default table or adjusting margins
                        for cell in row:
                            set_cell_margins(cell, top=20, bottom=20, left=40, right=40)
                    doc.add_paragraph()

                # Table
                if section.get("table"):
                    tbl_data = section["table"]
                    headers = tbl_data.get("headers", [])
                    rows = tbl_data.get("rows", [])
                    if headers and rows:
                        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
                        table.style = 'Table Grid'
                        table.autofit = True

                        hdr_cells = table.rows[0].cells
                        for i, header in enumerate(headers):
                            hdr_cells[i].text = ""
                            set_cell_bg(hdr_cells[i], header_hex)
                            set_cell_margins(hdr_cells[i], top=100, bottom=100, left=140, right=140)
                            para = hdr_cells[i].paragraphs[0]
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = para.add_run(str(header))
                            run.font.bold = True
                            run.font.size = Pt(style["body_size"])
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

                        for row_idx, row_data in enumerate(rows, 1):
                            row_cells = table.rows[row_idx].cells
                            for col_idx, value in enumerate(row_data):
                                if col_idx < len(row_cells):
                                    row_cells[col_idx].text = str(value)
                                    set_cell_margins(row_cells[col_idx], top=80, bottom=80, left=140, right=140)
                                    if row_idx % 2 == 0:
                                        set_cell_bg(row_cells[col_idx], _hex_from_rgb(style["table_alt_rgb"]))

                doc.add_paragraph()

            # === DAFTAR PUSTAKA ===
            references = data.get("references", [])
            if references:
                doc.add_paragraph()
                ref_h = doc.add_paragraph()
                ref_run = ref_h.add_run("Daftar Pustaka")
                ref_run.font.bold = True
                ref_run.font.size = Pt(style["heading_size"])
                ref_run.font.color.rgb = title_color

                accent_hex = _hex_from_rgb(style["accent_rgb"])
                for idx, ref in enumerate(references, 1):
                    if isinstance(ref, dict):
                        ref_text = ref.get("text", "")
                        ref_url = ref.get("url", "")
                    else:
                        ref_text = str(ref)
                        ref_url = ""

                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(4)
                    p.paragraph_format.left_indent = Inches(0.25)
                    main_run = p.add_run(f"{idx}. {ref_text}")
                    main_run.font.size = Pt(style["body_size"])

                    if ref_url:
                        p.add_run("  ").font.size = Pt(style["body_size"])
                        add_hyperlink(p, ref_url, ref_url, color_hex=accent_hex)

            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{data.get('filename', 'dokumen')}_{timestamp}.docx"
            filepath = OUTPUT_DIR / filename
            doc.save(filepath)

            return f"SUCCESS|{filepath}|Word ({style['label']}) berhasil dibuat: {filename}"
        except Exception as e:
            logger.error(f"[ADMIN] WordGenerator error: {e}", exc_info=True)
            return f"FAILED|{e}"


# ============================================================
# PDF Generator — sekarang support style + cover design + TOC
# ============================================================
class PDFGeneratorTool(BaseTool):
    name: str = "PDF Generator Tool"
    description: str = """Buat file PDF dengan style fleksibel & cover design.
    Input format JSON string:
    {
        "filename": "nama_file",
        "style": "formal" | "casual" | "creative" | "technical" | "educational",
        "cover": true,
        "title": "Judul Laporan",
        "subtitle": "Subjudul opsional",
        "author": "Nama Author",
        "toc": true,
        "sections": [
            {
                "heading": "Judul Bagian",
                "content": "Isi paragraf...",
                "list": ["Item 1", "Item 2"],
                "key_values": {"Nama": "Bima", "Jabatan": "Admin"},
                "image_path": "/path/to/image.png",
                "charts": [
                    {"type": "bar", "title": "Perbandingan Harga", "labels": ["Pinus","Jati","Mahoni"],
                     "datasets": [{"label": "Rp/m²", "data": [55000, 120000, 85000]}]}
                ],
                "table": {
                    "headers": ["Kolom1", "Kolom2"],
                    "rows": [["data1", "data2"]]
                }
            }
        ],
        "references": [
            {"text": "Kementerian PUPR (2025). Standar Harga Material Konstruksi. Jakarta.",
             "url": "https://pu.go.id"},
            {"text": "BPS (2026). Indeks Harga Bahan Bangunan.",
             "url": "https://bps.go.id"}
        ]
    }
    Field 'charts' (opsional, list per section): tipe 'bar'|'line'|'pie' — auto-render via matplotlib pakai warna dari 'style'.
    Field 'references' (opsional): daftar pustaka — auto render section terakhir dengan link clickable.
    URL referensi WAJIB valid (Wikipedia/situs resmi/jurnal open-access). JANGAN dikarang.
    Cocok untuk: laporan, proposal, resume, invoice, surat, tutorial, journal, certificate.
    Style menyesuaikan tone & color scheme."""

    def _safe_text(self, text: str) -> str:
        replacements = {
            '–': '-', '—': '-',
            '‘': "'", '’': "'",
            '“': '"', '”': '"',
            '…': '...', '•': '-',
            '→': '->', '←': '<-',
            '≥': '>=', '≤': '<=',
            '≠': '!=', '≈': '~=',
            ' ': ' ', '​': '',
        }
        for old, new in replacements.items():
            text = text.replace(old, new)
        return text.encode('latin-1', errors='replace').decode('latin-1')

    def _wrap_lines(self, pdf, text: str, max_width: float) -> int:
        """Estimasi jumlah baris untuk teks dengan lebar max_width (mm)."""
        if not text:
            return 1
        words = str(text).split()
        if not words:
            return 1
        usable = max_width - 4  # padding kiri-kanan
        lines = 1
        current = ""
        for word in words:
            candidate = (current + " " + word).strip() if current else word
            if pdf.get_string_width(candidate) <= usable:
                current = candidate
            else:
                lines += 1
                current = word
        return max(lines, 1)

    def _render_pdf_table(self, pdf, headers: list, rows: list, style: dict):
        """Render tabel dengan auto-wrap per cell (tanpa truncation), padding rapi.
        Setiap cell tingginya menyesuaikan konten terpanjang dalam row tersebut.
        """
        n_cols = len(headers)
        if n_cols == 0:
            return

        page_width = pdf.w - 2 * pdf.l_margin
        col_max_chars = []
        for c_idx in range(n_cols):
            header_len = len(str(headers[c_idx]))
            row_lens = [len(str(r[c_idx])) if c_idx < len(r) else 0 for r in rows]
            col_max_chars.append(max([header_len] + row_lens))
        total_chars = sum(col_max_chars) or 1
        col_widths = [max(page_width * (c / total_chars), 18) for c in col_max_chars]
        scale = page_width / sum(col_widths)
        col_widths = [w * scale for w in col_widths]

        line_h = 5.5

        pdf.set_font(style.get("pdf_font", "Helvetica"), "B", 9)
        pdf.set_fill_color(*style["table_header_rgb"])
        pdf.set_text_color(255, 255, 255)
        header_lines = max(self._wrap_lines(pdf, str(h), w) for h, w in zip(headers, col_widths))
        header_h = header_lines * line_h + 3

        if pdf.get_y() + header_h > pdf.page_break_trigger:
            pdf.add_page()

        x_start = pdf.l_margin
        y_start = pdf.get_y()
        for i, h in enumerate(headers):
            x_now = x_start + sum(col_widths[:i])
            pdf.set_xy(x_now, y_start)
            try:
                pdf.multi_cell(col_widths[i], line_h,
                               self._safe_text(str(h)),
                               border=1, fill=True, align="C",
                               max_line_height=line_h)
            except TypeError:
                # fpdf2 lama tidak support max_line_height
                pdf.multi_cell(col_widths[i], line_h,
                               self._safe_text(str(h)),
                               border=1, fill=True, align="C")
        pdf.set_xy(x_start, y_start + header_h)

        pdf.set_font(style.get("pdf_font", "Helvetica"), "", 9)
        pdf.set_text_color(30, 30, 30)
        for r_idx, row in enumerate(rows):
            row_padded = (list(row) + [""] * n_cols)[:n_cols]
            row_lines = max(self._wrap_lines(pdf, str(v), w) for v, w in zip(row_padded, col_widths))
            row_h = row_lines * line_h + 3

            if pdf.get_y() + row_h > pdf.page_break_trigger:
                pdf.add_page()

            if r_idx % 2 == 0:
                pdf.set_fill_color(*style["table_alt_rgb"])
            else:
                pdf.set_fill_color(255, 255, 255)

            y_row = pdf.get_y()
            for i, val in enumerate(row_padded):
                x_now = x_start + sum(col_widths[:i])
                pdf.set_xy(x_now, y_row)
                try:
                    pdf.multi_cell(col_widths[i], line_h,
                                   self._safe_text(str(val)),
                                   border=1, fill=True,
                                   max_line_height=line_h)
                except TypeError:
                    pdf.multi_cell(col_widths[i], line_h,
                                   self._safe_text(str(val)),
                                   border=1, fill=True)
            pdf.set_xy(x_start, y_row + row_h)

        pdf.ln(4)

    def _run(self, input_json: str) -> str:
        try:
            from fpdf import FPDF

            try:
                data = json.loads(input_json)
            except json.JSONDecodeError as e:
                return f"FAILED|JSON tidak valid: {e}"

            style_name = data.get("style", "formal")
            style = STYLES.get(style_name, STYLES["formal"])
            title_rgb = style["title_rgb"]
            accent_rgb = style["accent_rgb"]
            doc_title  = data.get("title", "Dokumen")
            doc_author = data.get("author", "B.I.M.A Core")

            # === Configurable typography (JSON overrides > style defaults) ===
            pdf_font = data.get("pdf_font", style.get("pdf_font", "Helvetica"))
            margins_cm = data.get("margins", style.get("margins_cm", {"top": 2.54, "bottom": 2.54, "left": 2.54, "right": 2.54}))
            do_justify = data.get("justify", style.get("justify", False))
            line_spacing_mult = data.get("line_spacing", style.get("line_spacing", 1.15))
            content_align = "J" if do_justify else "L"

            # === Mode Landscape untuk Sertifikat ===
            is_certificate = data.get("certificate", False)
            orientation = "L" if is_certificate else "P"

            # === Subclass FPDF: auto header & footer dengan nomor halaman ===
            _style_ref = style
            _title_ref = doc_title
            _author_ref = doc_author
            _accent_rgb = accent_rgb
            _title_rgb  = title_rgb
            _pdf_font   = pdf_font
            _body_start_page = [0]  # mutable agar bisa diupdate dari luar class

            def _to_roman_pdf(num):
                vals = [(1000,'m'),(900,'cm'),(500,'d'),(400,'cd'),(100,'c'),
                        (90,'xc'),(50,'l'),(40,'xl'),(10,'x'),(9,'ix'),(5,'v'),(4,'iv'),(1,'i')]
                result = ''
                for v, r in vals:
                    while num >= v:
                        result += r
                        num -= v
                return result

            class BimaFPDF(FPDF):
                def header(self):
                    # Hanya di halaman isi (bukan cover page 1)
                    if self.page_no() <= 1 and data.get("cover", True):
                        return
                    self.set_font(_pdf_font, "B", 8)
                    self.set_text_color(*_title_rgb)
                    self.cell(0, 8, self._safe_text_inner(_title_ref), new_x="LMARGIN", new_y="NEXT", align="L")
                    self.set_draw_color(*_accent_rgb)
                    self.set_line_width(0.3)
                    self.line(self.l_margin, self.get_y(), self.w - self.r_margin, self.get_y())
                    self.ln(3)

                def footer(self):
                    self.set_y(-14)
                    self.set_font(_pdf_font, "I", 8)
                    self.set_text_color(140, 140, 140)
                    pg = self.page_no()
                    # Roman/Arabic split untuk style academic
                    if _body_start_page[0] > 0 and pg >= _body_start_page[0]:
                        page_str = str(pg - _body_start_page[0] + 1)
                    elif _body_start_page[0] > 0:
                        page_str = _to_roman_pdf(pg)
                    else:
                        page_str = str(pg)
                    left_txt = f"{_author_ref}  -  {_style_ref['label']}"
                    self.cell(0, 8,
                              self._safe_text_inner(left_txt) + "     halaman " + page_str,
                              align="C")

                def _safe_text_inner(self, text):
                    reps = {'\u2013':'-','\u2014':'-','\u2018':"'",'\u2019':"'",'\u201c':'"','\u201d':'"','\u2026':'...'}
                    for old, new in reps.items():
                        text = text.replace(old, new)
                    return text.encode('latin-1', errors='replace').decode('latin-1')

            pdf = BimaFPDF(orientation=orientation)
            pdf.set_auto_page_break(auto=True, margin=margins_cm.get("bottom", 2.54) * 10)
            pdf.set_margins(
                left=margins_cm.get("left", 2.54) * 10,
                top=margins_cm.get("top", 2.54) * 10,
                right=margins_cm.get("right", 2.54) * 10
            )

            # === COVER PAGE (opsional) ===
            if data.get("cover", True):
                pdf.add_page()
                pdf.set_y(80)
                # Decorative bar
                pdf.set_fill_color(*accent_rgb)
                pdf.rect(10, 60, 30, 3, "F")

                pdf.set_font(pdf_font, "B", style["title_size"] + 8)
                pdf.set_text_color(*title_rgb)
                pdf.multi_cell(0, 12, self._safe_text(data.get("title", "Dokumen")), align="L")
                pdf.ln(4)

                if data.get("subtitle"):
                    pdf.set_font(pdf_font, "I", style["body_size"] + 4)
                    pdf.set_text_color(*accent_rgb)
                    pdf.multi_cell(0, 8, self._safe_text(data["subtitle"]), align="L")
                    pdf.ln(8)

                pdf.set_y(-50)
                pdf.set_font(pdf_font, "", 10)
                pdf.set_text_color(80, 80, 80)
                pdf.cell(0, 6, self._safe_text(f"Disusun oleh: {data.get('author', 'B.I.M.A Core')}"),
                         new_x="LMARGIN", new_y="NEXT")
                pdf.cell(0, 6, self._safe_text(f"Tanggal: {datetime.now().strftime('%d %B %Y')}"),
                         new_x="LMARGIN", new_y="NEXT")
                pdf.cell(0, 6, self._safe_text(f"Style: {style['label']}"),
                         new_x="LMARGIN", new_y="NEXT")

            # === ABSTRAK PDF (opsional) ===
            if data.get("abstract"):
                pdf.add_page()
                pdf.set_font(pdf_font, "B", style["heading_size"] + 2)
                pdf.set_text_color(*title_rgb)
                pdf.cell(0, 10, "ABSTRAK", new_x="LMARGIN", new_y="NEXT", align="C")
                pdf.ln(4)

                pdf.set_font(pdf_font, "", style["body_size"])
                pdf.set_text_color(30, 30, 30)
                # Abstrak single-spaced
                abs_lh = round(style["body_size"] * 1.15 * 0.39 + 5.5, 1)
                # Indent kiri-kanan 10mm ekstra
                old_l = pdf.l_margin
                old_r = pdf.r_margin
                pdf.set_left_margin(old_l + 10)
                pdf.set_right_margin(old_r + 10)
                pdf.multi_cell(0, abs_lh, self._safe_text(data["abstract"]), align=content_align)
                pdf.set_left_margin(old_l)
                pdf.set_right_margin(old_r)
                pdf.ln(4)

                # Keywords
                if data.get("keywords"):
                    pdf.set_left_margin(old_l + 10)
                    pdf.set_font(pdf_font, "B", style["body_size"])
                    kw_text = "Kata Kunci: "
                    pdf.cell(pdf.get_string_width(kw_text) + 2, 6, self._safe_text(kw_text))
                    pdf.set_font(pdf_font, "I", style["body_size"])
                    pdf.multi_cell(0, 6, self._safe_text(", ".join(data["keywords"])))
                    pdf.set_left_margin(old_l)

            # === HALAMAN ISI ===
            pdf.add_page()

            # Header tipis di halaman isi
            pdf.set_font(pdf_font, "B", style["title_size"])
            pdf.set_text_color(*title_rgb)
            pdf.cell(0, 12, self._safe_text(data.get("title", "Dokumen")),
                     new_x="LMARGIN", new_y="NEXT", align="C")

            pdf.set_font(pdf_font, "", 9)
            pdf.set_text_color(112, 112, 112)
            pdf.cell(0, 6, self._safe_text(
                f"{data.get('author', 'B.I.M.A Core')} | {datetime.now().strftime('%d %B %Y')}"
            ), new_x="LMARGIN", new_y="NEXT", align="C")
            pdf.ln(4)
            pdf.set_draw_color(*accent_rgb)
            pdf.set_line_width(0.5)
            pdf.line(pdf.l_margin, pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
            pdf.ln(8)

            # === Table of Contents (opsional, multi-level) ===
            sections = data.get("sections", [])
            if data.get("toc") and sections:
                pdf.set_font(pdf_font, "B", style["heading_size"] + 2)
                pdf.set_text_color(*title_rgb)
                pdf.cell(0, 10, "DAFTAR ISI", new_x="LMARGIN", new_y="NEXT", align="C")
                pdf.ln(4)
                pdf.set_text_color(40, 40, 40)
                for sec in sections:
                    if sec.get("heading"):
                        lvl = sec.get("level", 1)
                        indent = 5 * (lvl - 1)  # mm indent per level
                        if lvl == 1:
                            pdf.set_font(pdf_font, "B", style["body_size"])
                        else:
                            pdf.set_font(pdf_font, "", style["body_size"])
                        pdf.cell(indent, 7, "")
                        pdf.cell(0, 7, self._safe_text(sec["heading"]),
                                 new_x="LMARGIN", new_y="NEXT")
                pdf.ln(6)
                pdf.add_page()

            # === Mark body start page untuk Roman/Arabic split ===
            is_academic_pdf = (style_name == 'academic')
            if is_academic_pdf:
                _body_start_page[0] = pdf.page_no()

            # === SECTIONS (multi-level heading) ===
            body_lh = round(style["body_size"] * line_spacing_mult * 0.39 + 5.5, 1)  # line height proporsional
            for section in sections:
                if section.get("heading"):
                    lvl = section.get("level", 1)
                    lvl = max(1, min(lvl, 3))  # clamp 1-3

                    if lvl == 1:
                        pdf.set_font(pdf_font, "B", style["heading_size"] + 2)
                        pdf.set_text_color(*title_rgb)
                        pdf.multi_cell(0, 9, self._safe_text(section["heading"]))
                        # Garis tipis aksen di bawah heading level 1
                        pdf.set_draw_color(*accent_rgb)
                        pdf.set_line_width(0.4)
                        pdf.line(pdf.l_margin, pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
                        pdf.ln(3)
                    elif lvl == 2:
                        pdf.set_font(pdf_font, "B", style["heading_size"])
                        pdf.set_text_color(*title_rgb)
                        pdf.cell(3, 8, "")  # slight indent
                        pdf.multi_cell(0, 8, self._safe_text(section["heading"]))
                        pdf.ln(2)
                    elif lvl == 3:
                        pdf.set_font(pdf_font, "BI", style["heading_size"] - 1)
                        pdf.set_text_color(*title_rgb)
                        pdf.cell(6, 7, "")  # more indent
                        pdf.multi_cell(0, 7, self._safe_text(section["heading"]))
                        pdf.ln(2)

                if section.get("content"):
                    pdf.set_font(pdf_font, "", style["body_size"])
                    pdf.set_text_color(30, 30, 30)
                    pdf.multi_cell(0, body_lh, self._safe_text(section["content"]), align=content_align)
                    pdf.ln(3)

                # Bullet list
                if section.get("list"):
                    pdf.set_font(pdf_font, "", style["body_size"])
                    pdf.set_text_color(30, 30, 30)
                    for item in section["list"]:
                        pdf.cell(8, 6, "")
                        pdf.cell(4, 6, "-")
                        pdf.multi_cell(0, 6, self._safe_text(str(item)))
                    pdf.ln(2)

                # Key Values (untuk surat izin, detail rapi dengan titik dua sejajar)
                if section.get("key_values") and isinstance(section["key_values"], dict):
                    pdf.set_font(pdf_font, "", style["body_size"])
                    pdf.set_text_color(30, 30, 30)
                    for k, v in section["key_values"].items():
                        pdf.cell(40, 6, self._safe_text(str(k)))
                        pdf.cell(5, 6, ":")
                        pdf.multi_cell(0, 6, self._safe_text(str(v)))
                    pdf.ln(2)

                # Image embedding
                if section.get("image_path"):
                    img_path = Path(section["image_path"])
                    if img_path.exists() and img_path.is_file():
                        try:
                            pdf.image(str(img_path), w=160)
                            pdf.ln(4)
                        except Exception as img_err:
                            logger.warning(f"[ADMIN] Gagal embed image PDF {img_path}: {img_err}")

                # Chart rendering (matplotlib → embed PNG)
                for chart in section.get("charts", []) or []:
                    try:
                        chart_path = _render_chart(chart, style)
                        pdf.image(chart_path, w=170)
                        pdf.ln(4)
                    except Exception as chart_err:
                        logger.warning(f"[ADMIN] Gagal render chart PDF: {chart_err}")

                # Tables — auto wrap per cell, no truncation
                if section.get("table"):
                    tbl = section["table"]
                    headers = tbl.get("headers", [])
                    rows = tbl.get("rows", [])
                    if headers:
                        self._render_pdf_table(pdf, headers, rows, style)

            # === DAFTAR PUSTAKA (opsional) ===
            references = data.get("references", [])
            if references:
                pdf.ln(6)
                pdf.set_font(pdf_font, "B", style["heading_size"])
                pdf.set_text_color(*title_rgb)
                pdf.multi_cell(0, 9, self._safe_text("Daftar Pustaka"))
                pdf.ln(2)

                pdf.set_draw_color(*accent_rgb)
                pdf.set_line_width(0.3)
                pdf.line(pdf.l_margin, pdf.get_y(), 80, pdf.get_y())
                pdf.ln(4)

                for idx, ref in enumerate(references, 1):
                    if isinstance(ref, dict):
                        ref_text = ref.get("text", "")
                        ref_url = ref.get("url", "")
                    else:
                        ref_text = str(ref)
                        ref_url = ""

                    pdf.set_font(pdf_font, "", style["body_size"])
                    pdf.set_text_color(30, 30, 30)
                    pdf.multi_cell(0, 6, self._safe_text(f"{idx}. {ref_text}"))

                    if ref_url:
                        pdf.set_font(pdf_font, "U", style["body_size"] - 1)
                        pdf.set_text_color(*accent_rgb)
                        pdf.cell(8, 5, "")
                        pdf.cell(0, 5,
                                 self._safe_text(ref_url),
                                 link=ref_url,
                                 new_x="LMARGIN", new_y="NEXT")
                    pdf.ln(2)

            # Footer hanya untuk halaman terakhir jika subclass tidak di-override
            # (BimaFPDF sudah handle otomatis; ini fallback)

            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{data.get('filename', 'dokumen')}_{timestamp}.pdf"
            filepath = OUTPUT_DIR / filename
            pdf.output(str(filepath))

            return f"SUCCESS|{filepath}|PDF ({style['label']}) berhasil dibuat: {filename}"
        except ImportError:
            return "FAILED|fpdf2 belum terinstall. Jalankan: pip install fpdf2"
        except Exception as e:
            logger.error(f"[ADMIN] PDFGenerator error: {e}", exc_info=True)
            return f"FAILED|{e}"


# ============================================================
# Data Analysis Tool — dengan sinkronisasi warna dari STYLES
# ============================================================
class DataAnalysisTool(BaseTool):
    name: str = "Data Analysis & Chart Tool"
    description: str = """Baca file CSV/Excel, analisis data dengan pandas, dan hasilkan grafik visual.
    Input format: 'path_file|jenis_chart|kolom_x|kolom_y|style'
    Contoh: 'outputs/data.csv|bar|Bulan|Penjualan|formal'
    Field 'style' opsional (default formal) — warna chart menyesuaikan style preset BIMA.
    Jenis chart: bar, line, pie."""

    def _run(self, input_str: str) -> str:
        try:
            import matplotlib
            matplotlib.use('Agg')
            import matplotlib.pyplot as plt
            import pandas as pd

            if "|" not in input_str:
                return "FAILED|Format salah. Gunakan 'path_file|jenis_chart|kolom_x|kolom_y|style'"

            parts = input_str.split("|")
            filepath = parts[0].strip()
            chart_type = parts[1].strip().lower()

            # Style sync — ambil warna dari STYLES preset
            style_name = parts[4].strip().lower() if len(parts) > 4 else "formal"
            chart_style = STYLES.get(style_name, STYLES["formal"])
            primary_rgb = tuple(c / 255 for c in chart_style["accent_rgb"])
            header_rgb = tuple(c / 255 for c in chart_style["table_header_rgb"])
            palette = [
                primary_rgb,
                header_rgb,
                (0.40, 0.55, 0.85),
                (0.85, 0.55, 0.40),
                (0.55, 0.85, 0.55),
                (0.85, 0.45, 0.55),
            ]

            p = Path(filepath)
            if not p.exists():
                return f"FAILED|File tidak ditemukan: {filepath}"

            try:
                if p.suffix == '.csv':
                    df = pd.read_csv(filepath)
                elif p.suffix in ['.xlsx', '.xls']:
                    df = pd.read_excel(filepath)
                else:
                    return "FAILED|Format file harus CSV atau Excel."
            except Exception as read_err:
                return f"FAILED|Gagal baca file: {read_err}"

            fig, ax = plt.subplots(figsize=(10, 6), dpi=120)
            try:
                if chart_type == 'pie':
                    col_label = parts[2].strip()
                    col_val = parts[3].strip() if len(parts) > 3 and parts[3].strip() else None
                    if col_val:
                        chart_data = df.groupby(col_label)[col_val].sum()
                    else:
                        chart_data = df[col_label].value_counts()
                    ax.pie(chart_data.values, labels=chart_data.index,
                           autopct='%1.1f%%', startangle=140,
                           colors=palette[:len(chart_data)])
                    ax.set_title(f"Proporsi {col_val or col_label}",
                                fontsize=13, fontweight='bold',
                                color=tuple(c / 255 for c in chart_style["title_rgb"]))
                else:
                    x_col = parts[2].strip()
                    y_col = parts[3].strip() if len(parts) > 3 and parts[3].strip() else None
                    if not y_col:
                        return "FAILED|Kolom Y wajib diisi untuk chart bar/line."
                    df_agg = df.groupby(x_col)[y_col].sum().reset_index()
                    if chart_type == 'bar':
                        ax.bar(df_agg[x_col].astype(str), df_agg[y_col],
                               color=primary_rgb, edgecolor=header_rgb, linewidth=0.5)
                    elif chart_type == 'line':
                        ax.plot(df_agg[x_col].astype(str), df_agg[y_col],
                                marker='o', linewidth=2, color=primary_rgb,
                                markerfacecolor=header_rgb, markersize=6)
                        ax.grid(True, alpha=0.3)
                    ax.set_title(f"Total {y_col} per {x_col}",
                                fontsize=13, fontweight='bold',
                                color=tuple(c / 255 for c in chart_style["title_rgb"]))
                    ax.set_xlabel(x_col, fontsize=10)
                    ax.set_ylabel(y_col, fontsize=10)
                    ax.tick_params(axis='x', rotation=45)

                plt.tight_layout()
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                chart_filename = f"chart_{chart_type}_{timestamp}.png"
                chart_filepath = OUTPUT_DIR / chart_filename
                plt.savefig(chart_filepath, dpi=120, bbox_inches='tight', facecolor='white')
            finally:
                plt.close(fig)

            stats = df.describe().to_string()
            return f"SUCCESS|{chart_filepath}|\nStatistik:\n{stats}\n\nGrafik {chart_type} ({chart_style['label']}) disimpan di {chart_filepath}"

        except ImportError as e:
            return f"FAILED|Library kurang: {e}. Jalankan: pip install pandas matplotlib openpyxl"
        except Exception as e:
            logger.error(f"[ADMIN] DataAnalysis error: {e}", exc_info=True)
            return f"FAILED|Gagal analisis data: {e}"


# ============================================================
# Image Search Tool — cari + download gambar untuk embed di PDF/Word
# ============================================================
class ImageSearchTool(BaseTool):
    name: str = "Image Search Tool"
    description: str = """Cari + download gambar dari internet untuk disisipkan ke PDF/Word.
    PRIORITAS: Wikimedia Commons (license aman untuk jurnal/akademik) → fallback Serper Images.
    Cocok untuk: ilustrasi penelitian (anatomi tikus, mitosis sel), foto produk/material, gambar referensi desain.
    Input: keyword pencarian. Contoh: 'Rattus norvegicus laboratory', 'mitosis cell', 'kayu jati Jepara'.
    Output format: SUCCESS|/path/to/image.jpg|Sumber: ... | License: ... | URL: ..."""

    def _run(self, query: str) -> str:
        safe_q = "".join(c for c in query if c.isalnum() or c in (' ', '-', '_')).strip().replace(' ', '_')[:50]
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')

        # === 1. Wikimedia Commons (license aman utk akademik) ===
        try:
            resp = httpx.get(
                "https://commons.wikimedia.org/w/api.php",
                params={
                    "action": "query",
                    "format": "json",
                    "generator": "search",
                    "gsrsearch": query,
                    "gsrnamespace": 6,
                    "gsrlimit": 8,
                    "prop": "imageinfo",
                    "iiprop": "url|mime|extmetadata",
                    "iiurlwidth": 800,
                },
                headers={"User-Agent": "BIMA-Core-Research/1.0 (academic-use)"},
                timeout=15,
            )
            data = resp.json()
            pages = list(data.get("query", {}).get("pages", {}).values())
            pages.sort(key=lambda p: p.get("index", 999))

            for page in pages:
                info = (page.get("imageinfo") or [{}])[0]
                mime = info.get("mime", "")
                if mime not in ("image/jpeg", "image/png"):
                    continue
                img_url = info.get("thumburl") or info.get("url")
                if not img_url:
                    continue

                img_resp = httpx.get(
                    img_url,
                    headers={"User-Agent": "BIMA-Core-Research/1.0 (academic-use)"},
                    follow_redirects=True,
                    timeout=20,
                )
                if img_resp.status_code != 200 or len(img_resp.content) < 5000:
                    continue

                ext = ".jpg" if mime == "image/jpeg" else ".png"
                filename = f"img_{safe_q}_{timestamp}{ext}"
                filepath = OUTPUT_DIR / filename
                filepath.write_bytes(img_resp.content)

                title = page.get("title", "Unknown")
                meta = info.get("extmetadata", {})
                license_short = meta.get("LicenseShortName", {}).get("value", "Lihat sumber")
                artist = meta.get("Artist", {}).get("value", "")
                if "<" in artist:
                    import re
                    artist = re.sub(r"<[^>]+>", "", artist).strip()
                source_page = f"https://commons.wikimedia.org/wiki/{title.replace(' ', '_')}"

                return (
                    f"SUCCESS|{filepath}|"
                    f"Sumber: Wikimedia Commons - {title} | "
                    f"License: {license_short} | "
                    f"Artist: {artist or 'tidak disebutkan'} | "
                    f"URL Citation: {source_page}"
                )
            logger.info(f"[IMG] Wikimedia kosong untuk '{query}', fallback ke Serper")
        except Exception as e:
            logger.warning(f"[IMG] Wikimedia error: {e}, fallback ke Serper")

        # === 2. Serper Images (fallback) ===
        serper_key = os.environ.get("SERPER_API_KEY", "")
        if not serper_key:
            return f"FAILED|Tidak ada gambar di Wikimedia untuk '{query}', dan SERPER_API_KEY tidak tersedia untuk fallback"

        try:
            resp = httpx.post(
                "https://google.serper.dev/images",
                headers={"X-API-KEY": serper_key, "Content-Type": "application/json"},
                json={"q": query, "num": 8},
                timeout=15,
            )
            for item in resp.json().get("images", []):
                img_url = item.get("imageUrl") or ""
                lower_url = img_url.lower()
                if not (".jpg" in lower_url or ".jpeg" in lower_url or ".png" in lower_url):
                    continue

                try:
                    img_resp = httpx.get(
                        img_url,
                        headers={"User-Agent": "Mozilla/5.0 (BIMA-Core)"},
                        follow_redirects=True,
                        timeout=15,
                    )
                except Exception:
                    continue
                if img_resp.status_code != 200 or len(img_resp.content) < 5000:
                    continue

                ext = ".jpg" if (".jpg" in lower_url or ".jpeg" in lower_url) else ".png"
                filename = f"img_{safe_q}_{timestamp}{ext}"
                filepath = OUTPUT_DIR / filename
                filepath.write_bytes(img_resp.content)

                source = item.get("source", "Google Images")
                return (
                    f"SUCCESS|{filepath}|"
                    f"Sumber: {source} (via Google Images) | "
                    f"License: Periksa sumber asli — gunakan dengan hati-hati untuk publikasi akademik | "
                    f"URL Citation: {img_url}"
                )

            return f"FAILED|Tidak ada gambar valid ditemukan untuk '{query}' di Wikimedia maupun Serper"
        except Exception as e:
            logger.error(f"[IMG] Serper Images error: {e}", exc_info=True)
            return f"FAILED|Image search gagal: {e}"


# ============================================================
# Admin Agent — fleksibel, multi-style, multi-format
# ============================================================
admin_agent = Agent(
    role='Multi-Style Document Crafter',
    goal='Membuat dokumen Excel/Word/PDF dengan gaya tulisan dan layout yang menyesuaikan permintaan Bima — bukan hanya laporan formal.',
    backstory="""Kamu adalah Penulis Serbaguna B.I.M.A Core. Kamu bukan cuma tukang laporan formal.

    KAMU BISA BIKIN BANYAK TIPE DOKUMEN:
    - Laporan riset/bisnis (formal)
    - Skripsi / tesis / jurnal ilmiah (academic)
    - Blog post / newsletter (casual)
    - Cerita / esai naratif (creative)
    - Dokumentasi teknis / spesifikasi (technical)
    - Tutorial / panduan belajar (educational)
    - Resume/CV, invoice, surat formal, proposal proyek
    - Meeting minutes, project brief, journal harian
    - Certificate, recipe book, travel itinerary

    GAYA TULISAN ('style' di JSON tool):
    - "formal"       → Profesional, baku, struktur kaku (default untuk laporan/proposal)
    - "casual"       → Hangat, conversational, ramah (untuk blog/newsletter)
    - "creative"     → Ekspresif, naratif, deskriptif (untuk cerita/esai)
    - "technical"    → Presisi, padat, jargon teknis OK (untuk dokumentasi/spek)
    - "educational"  → Step-by-step, ramah pemula, contoh konkret (untuk tutorial)
    - "academic"     → Skripsi/tesis/jurnal ilmiah: Times New Roman 12pt, justify, spasi 1.5, margin 4-4-3-3 cm

    AUTO-DETECT STYLE AKADEMIK:
    - Kalau Bima bilang "skripsi", "tugas akhir", "tesis", "disertasi", "jurnal ilmiah",
      "makalah", "karya ilmiah" → otomatis pakai style "academic".
    - Style academic sudah preset: Times New Roman, justify, line spacing 1.5, margin kiri 4cm atas 4cm kanan 3cm bawah 3cm.

    TIPOGRAFI CONFIGURABLE (field opsional di JSON input — override preset style):
    - "font_family": "Arial" → override font Word (default dari style preset)
    - "pdf_font": "Helvetica" → override font PDF (pilihan: Helvetica, Times, Courier)
    - "margins": {"top": 4, "bottom": 3, "left": 4, "right": 3} → margin dalam CM
    - "line_spacing": 2.0 → override line spacing multiplier
    - "justify": true → rata kanan-kiri
    Contoh: kampus ITB minta margin 3-3-3-3 dan Arial 11pt → Bima bisa override lewat JSON tanpa ganti style.

    PILIHAN FORMAT:
    - Excel  → tabel, angka, perbandingan, rekap, formula (= di awal cell)
    - Word   → laporan naratif, surat, proposal, resume yang masih bisa diedit
    - PDF    → dokumen final siap kirim/cetak, support cover page + TOC

    GAMBAR/ILUSTRASI DI DOKUMEN:
    - Kalau Bima minta dokumen yang butuh gambar (jurnal penelitian, katalog, laporan visual),
      pakai ImageSearchTool DULU untuk cari + download gambar dari Wikimedia/Serper.
    - Tool ini return SUCCESS|/path/img.jpg|Sumber: ... | License: ...
    - Ambil path-nya, masukkan ke field "image_path" di section PDF/Word.
    - WAJIB sertakan info Sumber/License di "content" section sebagai caption (untuk integritas akademik).
    - Contoh workflow jurnal: search "Rattus norvegicus" → dapat path → embed di PDF section
      "Subjek Penelitian" dengan caption "Gambar 1. Rattus norvegicus. Sumber: Wikimedia Commons (CC BY-SA 4.0)".

    DIAGRAM/CHART/GRAFIK DI DOKUMEN:
    - PDFGeneratorTool & WordGeneratorTool sekarang punya field "charts" (list) per section.
    - Format: [{"type": "bar"|"line"|"pie", "title": "...", "labels": [...], "datasets": [{"label": "...", "data": [...]}]}]
    - Chart auto-render via matplotlib dengan warna sesuai 'style' preset.
    - Pakai untuk: perbandingan data, tren waktu, distribusi proporsi.
    - Untuk chart dari file CSV/Excel yang diupload Bima: pakai DataAnalysisTool DULU
      (tambahkan style di field ke-5: 'file.csv|bar|X|Y|academic' agar warna sinkron).
    - Untuk chart dari data JSON inline (hasil riset, hitungan manual): pakai field "charts" langsung.

    MULTI-LEVEL HEADING (SUB-BAB):
    - Setiap section sekarang support field "level": 1 | 2 | 3 (default 1).
    - Level 1 = BAB / Heading utama (bold, garis aksen di bawah)
    - Level 2 = Sub-bab (bold, indent sedikit)
    - Level 3 = Sub-sub-bab (bold italic, indent lebih)
    - Contoh JSON sections untuk skripsi:
      [{"heading": "BAB I PENDAHULUAN", "level": 1, "content": "..."},
       {"heading": "1.1 Latar Belakang", "level": 2, "content": "..."},
       {"heading": "1.1.1 Rumusan Masalah", "level": 3, "content": "..."},
       {"heading": "BAB II TINJAUAN PUSTAKA", "level": 1, "content": "..."}]
    - Word: menggunakan Heading 1/2/3 style (support TOC Word native)
    - PDF: ukuran font dan indent otomatis menyesuaikan level
    - TOC juga otomatis multi-level dengan indentasi per level.

    ABSTRAK & KATA KUNCI:
    - Field opsional di JSON root: "abstract": "teks abstrak...", "keywords": ["kata1", "kata2", "kata3"]
    - Auto-render halaman ABSTRAK terpisah (sebelum Daftar Isi):
      * Judul "ABSTRAK" centered, bold
      * Teks single-spaced (meskipun body 1.5), indent kiri-kanan
      * Keywords di bawah: "Kata Kunci: kata1, kata2, kata3" (bold label, italic value)
    - WAJIB diisi untuk style "academic" / dokumen skripsi.

    PENOMORAN HALAMAN ROMAN/ARABIC (otomatis untuk style academic):
    - Halaman depan (Cover, Abstrak, Daftar Isi) → angka Romawi kecil (i, ii, iii)
    - Halaman isi (BAB I dst) → angka Arab dimulai dari 1
    - Word: section break + pgNumType XML otomatis dihandle
    - PDF: footer otomatis switch format berdasarkan body_start_page
    - Untuk style selain academic: tetap pakai angka Arab biasa (backward compatible).

    SURAT FORMAL / SURAT IZIN (ATURAN KETAT):
    - Tata Letak Kontak (Header): Nama, alamat, dan nomor telepon JANGAN digabung dalam 1 baris. Buat secara vertikal di bagian awal konten (misal pakai multi-line paragraph).
    - Penerima Surat: Sebutkan nama instansi dan alamat perusahaan secara spesifik, jangan hanya menulis "di Tempat" agar lebih formal.
    - Data Diri: WAJIB gunakan field `key_values` agar titik dua (:) sejajar dan rapi. (Contoh: "key_values": {"Nama": "...", "Alamat": "..."})
    - Alasan & Durasi: Gunakan bahasa yang spesifik, logis, dan tidak ambigu (misal jika menunda atau izin sementara, sebutkan tanggalnya dengan jelas).
    - Tanda Baca: Pastikan ada spasi setelah koma dan titik. Jangan sampai menumpuk.

    INVOICE / TAGIHAN (ATURAN KETAT):
    - Header: Cantumkan tulisan "INVOICE" dengan jelas, Nomor Invoice unik, dan tanggal (Issue Date & Due Date).
    - Kontak: Info Pengirim dan Klien di bagian atas (bisa menggunakan `key_values`).
    - Rincian Item: WAJIB gunakan `table` dengan header: Deskripsi, Qty, Harga Satuan, dan Total.
    - Total: Tambahkan section untuk Subtotal, Pajak/Diskon, dan Total Amount Due secara jelas.

    MEETING MINUTES / NOTULEN RAPAT (ATURAN KETAT):
    - Header: Judul rapat, tanggal, waktu, lokasi.
    - Partisipan: Siapa yang hadir dan absen (gunakan list atau `key_values`).
    - Agenda & Diskusi: Jangan menulis transkrip kata demi kata. Fokus pada ringkasan objektif, keputusan (Decisions), dan Action Items (tugas, PIC, deadline). Gunakan list agar mudah dibaca.

    PROPOSAL PROYEK / BISNIS (ATURAN KETAT):
    - Wajib gunakan "cover": true dan "toc": true.
    - Struktur Standar: Executive Summary, Latar Belakang Masalah, Solusi/Metodologi, Timeline (jadwal kerja), Anggaran/RAB (wajib pakai tabel), dan Kesimpulan.
    
    TABEL & DATA KEY-VALUE DI DOKUMEN (ATURAN KETAT supaya rapi):
    - Untuk "surat izin", "biodata", atau data yang butuh tanda titik dua (:) sejajar, WAJIB gunakan field "key_values".
    - Untuk data berbentuk baris dan kolom yang banyak, gunakan "table".
    - Maksimal 5-6 kolom per tabel (lebih dari itu cell jadi sempit).
    - Cell value text idealnya <= 25 karakter. Kalau perlu lebih panjang, taruh di paragraf section, BUKAN tabel.
    - Tiap kata WAJIB dipisah spasi: "Rp 50.000" (bukan "Rp50000"), "Tahun 2026" (bukan "Tahun2026").
    - Header pakai Title Case singkat: "Nama Material", "Harga (Rp)", "Supplier".
    - Padding cell sudah otomatis dilebarkan (PDF auto-wrap, Word cell margin, Excel indent + min width 14).

    DAFTAR PUSTAKA / REFERENSI (WAJIB jika konten informatif/riset/edukatif/jurnal):
    - PDF & Word: pakai field "references" di JSON root (BUKAN di section):
      "references": [{"text": "Penulis (Tahun). Judul. Penerbit.", "url": "https://..."}, ...]
    - Excel: pakai field "references" di root → auto-bikin sheet "Referensi" dengan kolom No|Sumber|URL.
    - URL WAJIB valid & verifiable: Wikipedia, .gov, .edu, .org, jurnal open-access (DOI.org, arxiv.org),
      dokumentasi resmi vendor. JANGAN mengarang URL atau judul paper.
    - Kalau ragu link spesifiknya akurat → pakai homepage situsnya saja
      (mis. https://en.wikipedia.org daripada link artikel spesifik yang dikarang).
    - Link auto-clickable di PDF (FPDF link), Word (hyperlink XML), Excel (cell.hyperlink).

    ATURAN WAJIB:
    1. Deteksi gaya tulisan dari permintaan Bima — kalau dia bilang "skripsi"/"tugas akhir" → academic, "santai" → casual, "tutorial" → educational, dst
    2. SELALU sertakan "style" di JSON input tool. Default "formal" kalau gak yakin. Kalau konteks akademik → "academic".
    3. Konten HARUS substantial — minimal 3-5 sections untuk PDF/Word, minimal 5-10 baris untuk Excel
    4. Untuk PDF & Word: gunakan "cover": true dan "toc": true kalau dokumen laporan/berita >3 sections.
    5. Untuk dokumen akademik/jurnal dengan gambar: SELALU cantumkan caption sumber+license di bawah gambar
    6. Untuk dokumen riset/laporan/tutorial: WAJIB tambahkan field "references" — minimal 3 sumber valid.
    7. WAJIB return: SUCCESS|path_file|keterangan ATAU FAILED|alasan
    8. Untuk style "academic": JANGAN override font/margin/spacing kecuali Bima eksplisit minta.
    9. Untuk style "academic": WAJIB isi "abstract" dan "keywords" di JSON input.
    10. Untuk skripsi: GUNAKAN "level" di sections (1 untuk BAB, 2 untuk sub-bab, 3 untuk sub-sub-bab).

    Output kamu siap dikirim ke Bima — tidak perlu diedit lagi.""",
    llm=admin_llm,
    tools=[ExcelGeneratorTool(), WordGeneratorTool(), PDFGeneratorTool(), DataAnalysisTool(), ImageSearchTool()],
    allow_delegation=True,
    verbose=True
)
